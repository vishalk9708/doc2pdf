# from docx import Document
# import os

# def main():
#     template_file_path = 'investordoc.docx'
#     output_file_path = 'result.docx'

#     variables = {
#         "${InvestorName}": "Parag Poddar",
#     }

#     template_document = Document(template_file_path)

#     for variable_key, variable_value in variables.items():
#         for paragraph in template_document.paragraphs:
#             replace_text_in_paragraph(paragraph, variable_key, variable_value)

#         for table in template_document.tables:
#             for col in table.columns:
#                 for cell in col.cells:
#                     for paragraph in cell.paragraphs:
#                         replace_text_in_paragraph(paragraph, variable_key, variable_value)

#     template_document.save(output_file_path)


# def replace_text_in_paragraph(paragraph, key, value):
#     if key in paragraph.text:
#         inline = paragraph.runs
#         for item in inline:
#             if key in item.text:
#                 item.text = item.text.replace(key, value)


# if __name__ == '__main__':
#     main()
from docx import Document
from docx2pdf import convert
from simple_colors import *
import json 
doc = Document('investordoc.docx')
f = open('replacementData.json')
  
# returns JSON object as 
# a dictionary
replacements = json.load(f)

for paragraph in doc.paragraphs:
    for key in replacements:
        paragraph.text = paragraph.text.replace(key, replacements[key])

doc.save('abcde.docx')
convert('abcde.docx', 'output.pdf')